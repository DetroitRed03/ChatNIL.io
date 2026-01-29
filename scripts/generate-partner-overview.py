#!/usr/bin/env python3
"""
Generate ChatNIL Partner Overview Document
Professional Word document for stakeholders, investors, and school administrators
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ChatNIL brand color
CHATNIL_ORANGE = RGBColor(249, 115, 22)  # #F97316
DARK_GRAY = RGBColor(31, 41, 55)  # #1F2937
LIGHT_GRAY = RGBColor(107, 114, 128)  # #6B7280

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_heading(doc, text, level=1):
    """Create a styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = CHATNIL_ORANGE
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = DARK_GRAY
            run.font.size = Pt(18)
        elif level == 3:
            run.font.color.rgb = DARK_GRAY
            run.font.size = Pt(14)
    return heading

def add_bullet_list(doc, items, bold_first_part=False):
    """Add a bulleted list"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_first_part and ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(item)

def add_table(doc, headers, rows, first_col_bold=False):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], 'F97316')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = str(cell_text)
            if first_col_bold and col_idx == 0:
                for paragraph in row.cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            # Alternating row colors
            if row_idx % 2 == 0:
                set_cell_shading(row.cells[col_idx], 'FFF7ED')

    doc.add_paragraph()  # Space after table
    return table

def create_document():
    doc = Document()

    # ==================== COVER PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[ChatNIL Logo]')
    run.font.size = Pt(14)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Platform Overview')
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Compliance-First Approach to NIL Education')
    run.font.size = Pt(24)
    run.font.color.rgb = CHATNIL_ORANGE

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('January 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_paragraph()

    # Confidential notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONFIDENTIAL')
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_GRAY

    add_page_break(doc)

    # ==================== TABLE OF CONTENTS ====================
    create_heading(doc, 'Table of Contents', 1)

    toc_items = [
        ('Executive Summary', '3'),
        ('Section 1: The Problem We Solve', '4'),
        ('Section 2: High School Student Experience', '6'),
        ('Section 3: College Athlete Experience', '9'),
        ('Section 4: Parent Experience', '12'),
        ('Section 5: Compliance Officer Experience', '14'),
        ('Section 6: The 6-Dimension Scoring System', '18'),
        ('Section 7: Why ChatNIL?', '20'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run('\t' * 6)
        p.add_run(page)

    add_page_break(doc)

    # ==================== EXECUTIVE SUMMARY ====================
    create_heading(doc, 'Executive Summary', 1)

    p = doc.add_paragraph()
    run = p.add_run('The NIL landscape is broken.')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        'Two competing frameworks—the SCORE Act and the House Settlement—have created massive confusion '
        'about what constitutes legitimate third-party NIL versus disguised pay-for-play. Schools need '
        'compliance tools NOW, not after the dust settles.'
    )

    doc.add_paragraph(
        'ChatNIL doesn\'t try to solve pay-for-play. Instead, we clearly define, document, and enforce '
        'what legitimate third-party NIL looks like. We are the neutral compliance authority, not a '
        'marketplace participant.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Our Position: ')
    run.bold = True
    p.add_run('Neutral compliance authority, not marketplace participant.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Four User Types:')
    run.bold = True
    run.font.size = Pt(12)

    add_bullet_list(doc, [
        'High School Students: Education & preparation for NIL',
        'College Athletes: Compliance validation & deal scoring',
        'Parents: Oversight, consent, and peace of mind',
        'Compliance Officers: Institutional management & NCAA documentation'
    ], bold_first_part=True)

    p = doc.add_paragraph()
    run = p.add_run('What Makes Us Different:')
    run.bold = True

    doc.add_paragraph(
        'We don\'t connect athletes to brands. We don\'t take a cut of deals. We don\'t compete with '
        'collectives or agencies. We are the referee, not a player. This neutrality is why schools trust us '
        'and why our compliance scoring carries weight.'
    )

    add_page_break(doc)

    # ==================== SECTION 1: THE PROBLEM WE SOLVE ====================
    create_heading(doc, 'Section 1: The Problem We Solve', 1)

    create_heading(doc, 'The Current NIL Mess', 2)

    doc.add_paragraph(
        'The NIL landscape is governed by two conflicting frameworks that have created unprecedented '
        'confusion for athletes, schools, and brands:'
    )

    add_table(doc,
        ['Framework', 'Source', 'Key Feature'],
        [
            ['SCORE Act', 'Federal/Government', 'Government standards for NIL activities'],
            ['House Settlement', 'NCAA', '~$20.5M salary cap framework per school'],
        ]
    )

    create_heading(doc, 'The Confusion', 2)

    add_bullet_list(doc, [
        'Pay-for-play NIL: Schools pay athletes directly (capped under House Settlement)',
        'Third-party NIL: Brands pay athletes for endorsements (uncapped, legitimate)',
        'Money flows between these pots with no clear boundary',
        'Schools and collectives mask pay-for-play as third-party NIL',
        'No one knows what\'s allowed anymore'
    ])

    create_heading(doc, 'Why This Matters', 2)

    add_table(doc,
        ['Stakeholder', 'Risk'],
        [
            ['Athletes', 'Risk losing eligibility for unknowing violations'],
            ['Schools', 'Risk NCAA sanctions and investigation'],
            ['Brands', 'Risk association with compliance violations'],
            ['Parents', 'Don\'t know what deals are safe for their child'],
        ]
    )

    create_heading(doc, 'ChatNIL\'s Answer', 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"We don\'t solve pay-for-play. We define, document, and enforce what legitimate third-party NIL looks like."')
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = CHATNIL_ORANGE

    doc.add_paragraph()
    doc.add_paragraph(
        'Our 6-dimension scoring system creates a clear, auditable standard for what constitutes '
        'legitimate third-party NIL. Every deal is scored, documented, and defensible.'
    )

    add_page_break(doc)

    # ==================== SECTION 2: HIGH SCHOOL STUDENT EXPERIENCE ====================
    create_heading(doc, 'Section 2: High School Student Experience', 1)

    create_heading(doc, 'WHAT', 2)

    doc.add_paragraph(
        'The High School Student dashboard is an education-focused experience that prepares young '
        'athletes for NIL BEFORE they get to college. We don\'t help them sign deals—most states '
        'restrict or prohibit that anyway. Instead, we teach them the knowledge they\'ll need when '
        'the time comes.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Key Components:')
    run.bold = True

    add_bullet_list(doc, [
        'Discovery Through Conversation: AI-guided learning that asks questions first',
        '4-Pillar Learning Path: Identity, Business, Money, Legacy',
        'State Rules Education: What\'s allowed in their specific state',
        'Parent Consent Integration: Legal requirement, built in from day one',
        'Badge & Streak Gamification: Motivation to keep learning'
    ])

    create_heading(doc, 'Dashboard Elements', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Journey Progress', 'Shows current pillar and completion percentage'],
            ['Continue Conversation', 'Primary CTA - resumes AI-guided discovery'],
            ['State Rules Card', 'Shows state-specific HS NIL rules'],
            ['Parent Consent Status', 'Shows if parent has approved'],
            ['Chapters Grid', '4 pillars with lock/unlock status'],
            ['Badge Collection', 'Educational badges earned'],
            ['Streak Tracker', 'Daily engagement motivation'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'WHY', 2)

    p = doc.add_paragraph()
    run = p.add_run('Why Education First?')
    run.bold = True

    add_bullet_list(doc, [
        'Most states restrict or prohibit HS NIL deals',
        'Athletes need to understand rules BEFORE signing anything',
        'Building knowledge foundation prevents future mistakes',
        'Parents need assurance this is educational, not transactional'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why Discovery Through Conversation?')
    run.bold = True

    add_bullet_list(doc, [
        'Meets students where they are (conversational, not lecture)',
        'AI asks questions first (not waiting for student to know what to ask)',
        'Collects profile data while teaching (efficient)',
        'Unlocks chapters through engagement (gamified progression)'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why 4 Pillars?')
    run.bold = True

    add_table(doc,
        ['Pillar', 'Focus', 'Why It Matters'],
        [
            ['Identity', 'Know yourself', 'Before selling yourself, understand what makes you unique'],
            ['Business', 'Understand the rules', 'Learn the game before playing it'],
            ['Money', 'Financial literacy', 'Prevents exploitation and surprise tax bills'],
            ['Legacy', 'Think long-term', 'NIL should build toward something bigger'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'HOW', 2)

    p = doc.add_paragraph()
    run = p.add_run('How Discovery Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Student logs in → AI Coach initiates conversation',
        'AI asks about sport, goals, social media presence',
        'Student answers naturally → System extracts data',
        'After 5 days of conversation → Chapter unlocks',
        'Student can take quiz to earn badges',
        'Progression: Identity → Business → Money → Legacy'
    ])

    p = doc.add_paragraph()
    run = p.add_run('How Parent Consent Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Student signs up → Enters parent email',
        'Parent receives consent request email',
        'Parent clicks link → Creates account or logs in',
        'Parent reviews → Approves or denies',
        'If approved → Student can proceed',
        'If denied → Student sees "Parent did not approve"'
    ])

    p = doc.add_paragraph()
    run = p.add_run('What They DON\'T See:')
    run.bold = True

    add_bullet_list(doc, [
        'No deal validation (they\'re not signing deals)',
        'No compliance scoring (not relevant yet)',
        'No brand matching (we\'re not a marketplace for them)',
        'No messaging (no one to message)'
    ])

    add_page_break(doc)

    # ==================== SECTION 3: COLLEGE ATHLETE EXPERIENCE ====================
    create_heading(doc, 'Section 3: College Athlete Experience', 1)

    create_heading(doc, 'WHAT', 2)

    doc.add_paragraph(
        'The College Athlete dashboard is a compliance-focused experience that helps athletes validate '
        'deals and stay eligible. Unlike marketplace platforms, we don\'t connect them to brands—we help '
        'them ensure the deals they find are legitimate and compliant.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Key Components:')
    run.bold = True

    add_bullet_list(doc, [
        'Compliance Status Overview: GREEN/YELLOW/RED at a glance',
        'Deal Validator: 6-dimension scoring system',
        'Active Deals List: All deals sorted by compliance severity',
        'Tax Tracker: YTD earnings and estimated tax obligations',
        'State Rules Reference: State-specific NIL regulations'
    ])

    create_heading(doc, 'Dashboard Elements', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Compliance Status Banner', 'Overall GREEN/YELLOW/RED status'],
            ['Validate New Deal', 'Primary CTA - opens validation wizard'],
            ['Deals List', 'All deals sorted by compliance severity'],
            ['Tax Tracker', 'YTD earnings and estimated tax'],
            ['State Rules', 'State-specific NIL regulations'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'WHY', 2)

    p = doc.add_paragraph()
    run = p.add_run('Why Compliance-Focused (Not Marketplace)?')
    run.bold = True

    add_bullet_list(doc, [
        'Marketplace puts us INSIDE the confusion',
        'Compliance makes us the NEUTRAL ARBITER',
        'Schools will pay for compliance tools',
        'Athletes trust a validator more than a matchmaker'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why 6-Dimension Scoring?')
    run.bold = True

    doc.add_paragraph('This is our core patent. Each dimension answers a specific question:')

    add_table(doc,
        ['Dimension', 'Weight', 'Question It Answers'],
        [
            ['Policy Fit', '30%', 'Does this comply with NCAA rules and state law?'],
            ['Document Hygiene', '20%', 'Is there a clean contract without red flags?'],
            ['FMV Verification', '15%', 'Is the payment market-rate or suspiciously inflated?'],
            ['Tax Readiness', '15%', 'Does the athlete understand their tax obligations?'],
            ['Brand Safety', '10%', 'Is this an appropriate brand category?'],
            ['Guardian Consent', '10%', 'If minor, has parent approved?'],
        ],
        first_col_bold=True
    )

    p = doc.add_paragraph()
    run = p.add_run('Score Thresholds:')
    run.bold = True

    add_bullet_list(doc, [
        '🟢 GREEN (80-100): Proceed with confidence',
        '🟡 YELLOW (50-79): Issues exist but fixable',
        '🔴 RED (0-49): Do not proceed - serious compliance risk'
    ])

    create_heading(doc, 'HOW', 2)

    p = doc.add_paragraph()
    run = p.add_run('How Deal Validation Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Athlete clicks "Validate New Deal"',
        'Step 1: Enter deal basics (who, what, how much)',
        'Step 2: Answer compliance questions (booster? performance-based?)',
        'Step 3: See compliance score with dimension breakdown',
        'If GREEN → Save and proceed',
        'If YELLOW → See specific issues and fix recommendations',
        'If RED → Do not proceed, serious compliance risk'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Pay-for-Play Red Flags (Auto-Detected):')
    run.bold = True

    add_bullet_list(doc, [
        'Compensation >2x fair market value',
        'Booster or collective involvement',
        'Payment tied to athletic performance (touchdowns, wins)',
        'School or athletic department connection',
        'No clear deliverables or vague requirements'
    ])

    p = doc.add_paragraph()
    run = p.add_run('What They DON\'T See:')
    run.bold = True

    add_bullet_list(doc, [
        'No brand discovery (we don\'t connect them to brands)',
        'No agency matching (we don\'t play matchmaker)',
        'No campaign invites (no marketplace)',
        'No messaging (no one to message)'
    ])

    add_page_break(doc)

    # ==================== SECTION 4: PARENT EXPERIENCE ====================
    create_heading(doc, 'Section 4: Parent Experience', 1)

    create_heading(doc, 'WHAT', 2)

    doc.add_paragraph(
        'The Parent dashboard provides read-only oversight of their child\'s NIL education journey. '
        'Parents can monitor progress, manage consent, and receive notifications—but they don\'t '
        'control the content or make decisions for their child.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Key Components:')
    run.bold = True

    add_bullet_list(doc, [
        'Child Progress Overview: Visual progress tracking',
        'Consent Management: Approve, revoke, or modify consent',
        'Activity Feed: Recent child activities',
        'Notification Settings: Email preferences'
    ])

    create_heading(doc, 'Dashboard Elements', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Child Card', 'Shows child\'s name, school, sport, progress'],
            ['Learning Progress Bar', 'Visual completion percentage'],
            ['Current Chapter', 'Which pillar child is working on'],
            ['Consent Status', 'Approved/Pending/Denied with management'],
            ['Activity Feed', 'Recent child activities'],
            ['Notification Settings', 'Email preferences'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'WHY', 2)

    p = doc.add_paragraph()
    run = p.add_run('Why Read-Only?')
    run.bold = True

    add_bullet_list(doc, [
        'Parents oversee, they don\'t control',
        'Builds trust without helicopter parenting',
        'Child owns their learning journey',
        'Legal requirement for consent, not content control'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why Activity Feed?')
    run.bold = True

    add_bullet_list(doc, [
        'Parents want to know their child is engaged',
        'Shows badges earned, quizzes completed',
        'Builds confidence platform is educational',
        'No need to ask child "what did you learn?"'
    ])

    create_heading(doc, 'HOW', 2)

    p = doc.add_paragraph()
    run = p.add_run('How Consent Flow Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Child signs up → System requires parent email',
        'Parent receives email: "[Child] wants to join ChatNIL"',
        'Email explains: What ChatNIL is, what child will learn, what we DON\'T do',
        'Parent clicks "Approve" → Creates account, consent recorded',
        'Parent can monitor progress from their dashboard',
        'Parent can revoke consent at any time'
    ])

    p = doc.add_paragraph()
    run = p.add_run('What They DON\'T See:')
    run.bold = True

    add_bullet_list(doc, [
        'Child\'s conversation content (privacy)',
        'Ability to edit child\'s profile',
        'Ability to submit deals on child\'s behalf',
        'Any marketplace or deal features'
    ])

    add_page_break(doc)

    # ==================== SECTION 5: COMPLIANCE OFFICER EXPERIENCE ====================
    create_heading(doc, 'Section 5: Compliance Officer Experience', 1)

    create_heading(doc, 'WHAT', 2)

    doc.add_paragraph(
        'The Compliance Officer dashboard provides institutional oversight of all athletes at their '
        'school or organization. It\'s designed for efficiency at scale—finding problems quickly, '
        'not browsing paperwork.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Three-Level Navigation:')
    run.bold = True

    add_bullet_list(doc, [
        'Level 1 - Overview Dashboard: Aggregate stats, alerts, deadlines',
        'Level 2 - Athlete List: Paginated, searchable, filterable',
        'Level 3 - Athlete Detail: Individual history, deals, overrides'
    ])

    create_heading(doc, 'Level 1: Overview Dashboard', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Needs Attention List', 'Athletes with RED/YELLOW status'],
            ['Deadline Tracker', 'NCAA reporting deadlines (5-day rule)'],
            ['Compliance Stats', 'GREEN/YELLOW/RED/No Deals counts'],
            ['Sport Breakdown', 'Compliance by sport'],
            ['Quick Actions', 'Search, Roster, Export buttons'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'Level 2: Athlete List', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Search', 'Find athletes by name or ID'],
            ['Filters', 'Status, sport, deal count'],
            ['Paginated Table', 'Handle 1000+ athletes efficiently'],
            ['Bulk Actions', 'Mark reviewed, export, message'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'Level 3: Athlete Detail', 3)

    add_table(doc,
        ['Component', 'Purpose'],
        [
            ['Compliance Summary', 'Overall status and risk level'],
            ['Deals List', 'All deals with scores and issues'],
            ['Override Panel', 'Manual score adjustment with audit'],
            ['Audit Trail', 'Complete action history'],
        ],
        first_col_bold=True
    )

    create_heading(doc, 'WHY', 2)

    p = doc.add_paragraph()
    run = p.add_run('Why Three Levels?')
    run.bold = True

    add_bullet_list(doc, [
        'Compliance officers don\'t browse, they find problems',
        'Overview shows what needs attention NOW',
        'List lets them filter to specific concerns',
        'Detail lets them take action on individuals'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why "Needs Attention" First?')
    run.bold = True

    add_bullet_list(doc, [
        '1,000 athletes, maybe 50 have issues',
        'Don\'t waste time on compliant athletes',
        'Surface problems, not paperwork',
        'RED first, then YELLOW, then GREEN'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why Deadline Tracker?')
    run.bold = True

    add_bullet_list(doc, [
        'NCAA requires deal disclosure within 5 business days',
        'Missing deadlines = NCAA violation',
        'Proactive alerts prevent compliance failures',
        'Shows deals due in 2 days vs 5 days'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Why Override Capability?')
    run.bold = True

    add_bullet_list(doc, [
        'Algorithms aren\'t perfect',
        'Compliance officer may have information system doesn\'t',
        'Example: "Booster Collective" name triggers flag, but officer verified it\'s unaffiliated',
        'All overrides logged for audit trail'
    ])

    create_heading(doc, 'HOW', 2)

    p = doc.add_paragraph()
    run = p.add_run('How Scale is Handled:')
    run.bold = True

    add_bullet_list(doc, [
        'Server-side pagination (never load 1000+ records)',
        'Server-side filtering (database does the work)',
        'Server-side search (fast text search)',
        'Cached aggregates (overview stats refresh every 5 min)'
    ])

    p = doc.add_paragraph()
    run = p.add_run('How Override Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Officer views athlete detail',
        'Selects deal to override',
        'Chooses new status (can only improve, not worsen)',
        'Enters required reason (min 50 characters)',
        'System records override with officer ID and timestamp',
        'Audit trail shows: original score → new score + reason'
    ])

    p = doc.add_paragraph()
    run = p.add_run('How NCAA Export Works:')
    run.bold = True

    add_bullet_list(doc, [
        'Officer clicks "Generate NCAA Report"',
        'Selects date range and filters',
        'System generates CSV with required fields',
        'Download for submission to NCAA'
    ])

    doc.add_paragraph('Export includes: Athlete name, sport, ID, deal details, third party info, amount, dates, compliance status, all six dimension scores.')

    p = doc.add_paragraph()
    run = p.add_run('What They DON\'T See:')
    run.bold = True

    add_bullet_list(doc, [
        'Athletes at other institutions (data isolation)',
        'Ability to edit athlete profiles',
        'Ability to delete history',
        'Conversation content (athlete privacy)',
        'Marketplace features'
    ])

    add_page_break(doc)

    # ==================== SECTION 6: THE 6-DIMENSION SCORING SYSTEM ====================
    create_heading(doc, 'Section 6: The 6-Dimension Scoring System', 1)

    p = doc.add_paragraph()
    run = p.add_run('The Core Patent')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = CHATNIL_ORANGE

    doc.add_paragraph(
        'Our 6-dimension scoring system answers one critical question: "Is this deal legitimate '
        'third-party NIL or disguised pay-for-play?" Each dimension evaluates a specific aspect '
        'of deal legitimacy.'
    )

    create_heading(doc, 'Dimension 1: Policy Fit (30%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'NCAA rules compliance',
        'State law compliance',
        'School-specific policies',
        'Booster/collective involvement flags'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '100: Fully compliant with all regulations',
        '-40: School-affiliated deal',
        '-50: Booster-connected deal',
        '0: Performance-based compensation (auto-fail)'
    ])

    create_heading(doc, 'Dimension 2: Document Hygiene (20%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'Contract present?',
        'Prohibited terms?',
        'Clear deliverables?',
        'Defined duration?'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '100: Clean contract with all elements',
        '-30: No contract provided',
        '-30: Prohibited term found (per term)',
        '-20: Vague deliverables',
        '-10: No duration specified'
    ])

    create_heading(doc, 'Dimension 3: FMV Verification (15%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'Is payment reasonable for this athlete\'s reach?',
        'Compared to market benchmarks',
        'Variance from expected value'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '95-100: Within market range',
        '75: 50% above market (minor concern)',
        '50: 100% above market (significant)',
        '20: 200%+ above market (major red flag)'
    ])

    create_heading(doc, 'Dimension 4: Tax Readiness (15%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'Has athlete acknowledged tax obligations?',
        'Will they receive 1099?',
        'Quarterly payment awareness'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '100: Tax obligations acknowledged',
        '-40: Not acknowledged',
        'Additional reminders based on amount'
    ])

    create_heading(doc, 'Dimension 5: Brand Safety (10%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'Prohibited categories (alcohol, tobacco, gambling, etc.)',
        'Caution categories (supplements, crypto, etc.)',
        'Brand verification'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '0: Prohibited category (auto-fail)',
        '-20: Caution category',
        '-15: Unverified third party'
    ])

    create_heading(doc, 'Dimension 6: Guardian Consent (10%)', 2)

    p = doc.add_paragraph()
    run = p.add_run('What It Checks:')
    run.bold = True

    add_bullet_list(doc, [
        'Is athlete a minor?',
        'Has parent/guardian approved?'
    ])

    p = doc.add_paragraph()
    run = p.add_run('Scoring Logic:')
    run.bold = True

    add_bullet_list(doc, [
        '100: Adult (N/A) or consent approved',
        '40: Consent pending',
        '0: Consent denied or missing'
    ])

    create_heading(doc, 'Combined Score Thresholds', 2)

    add_table(doc,
        ['Score Range', 'Status', 'Meaning'],
        [
            ['80-100', '🟢 GREEN', 'Legitimate third-party NIL - Proceed with confidence'],
            ['50-79', '🟡 YELLOW', 'Concerns to address - Issues exist but fixable'],
            ['0-49', '🔴 RED', 'Likely pay-for-play or serious violation - Do not proceed'],
        ],
        first_col_bold=True
    )

    add_page_break(doc)

    # ==================== SECTION 7: WHY CHATNIL? ====================
    create_heading(doc, 'Section 7: Why ChatNIL?', 1)

    create_heading(doc, 'For Schools', 2)

    add_bullet_list(doc, [
        'Compliance tooling they need NOW, not after regulations settle',
        'Defensible documentation for NCAA audits and investigations',
        'Proactive problem identification before violations occur',
        'Scales to thousands of athletes without additional staff'
    ])

    create_heading(doc, 'For Athletes', 2)

    add_bullet_list(doc, [
        'Know their deals are clean before signing',
        'Protect their eligibility with documented compliance',
        'Understand their tax and legal obligations',
        'Preparation before college (HS students)'
    ])

    create_heading(doc, 'For Parents', 2)

    add_bullet_list(doc, [
        'Assurance the platform is educational, not transactional',
        'Visibility into child\'s learning progress',
        'Control via consent management',
        'Trust in a platform that prioritizes their child\'s future'
    ])

    create_heading(doc, 'For the NCAA/Government', 2)

    add_bullet_list(doc, [
        'Clear third-party NIL documentation standards',
        'Auditable compliance records for investigation',
        'Neutral enforcement of standards (not a marketplace)',
        'Supports legitimate NIL while flagging disguised pay-for-play'
    ])

    create_heading(doc, 'Our Competitive Advantage', 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"We\'re the referee, not a player."')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = CHATNIL_ORANGE

    doc.add_paragraph()

    add_bullet_list(doc, [
        'We\'re not trying to make money on deals',
        'We\'re the referee, not a player in the NIL marketplace',
        'Schools trust us because we\'re not conflicted',
        'Athletes trust us because we protect them, not profit from them',
        'Our compliance scoring carries weight because we\'re neutral'
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('ChatNIL: Compliance-First NIL Education')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = CHATNIL_ORANGE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    output_dir = '/Users/verrelbricejr./ChatNIL.io/docs'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'ChatNIL_Platform_Overview.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
