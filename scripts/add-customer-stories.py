#!/usr/bin/env python3
"""
Add Customer Stories section to ChatNIL Partner Overview Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ChatNIL brand color
CHATNIL_ORANGE = RGBColor(249, 115, 22)  # #F97316
DARK_GRAY = RGBColor(31, 41, 55)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_customer_stories(doc):
    """Add Section 8: Customer Stories to the document"""

    # Page break before new section
    doc.add_page_break()

    # Section Header
    heading = doc.add_heading('Section 8: Customer Stories', 1)
    for run in heading.runs:
        run.font.color.rgb = CHATNIL_ORANGE
        run.font.size = Pt(24)

    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run('Real Problems, Real Solutions')
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GRAY
    run.italic = True

    doc.add_paragraph()

    # Intro paragraph
    doc.add_paragraph(
        'These stories represent the real challenges our four user types face in the NIL landscape—and '
        'how ChatNIL\'s compliance-first approach solves them. Each persona is fictional, but the problems '
        'they face are happening to thousands of athletes, parents, and compliance officers right now.'
    )

    doc.add_paragraph()

    # ==================== JASMINE'S STORY ====================
    heading = doc.add_heading("Jasmine's Story: \"I Almost Signed the Wrong Deal\"", 2)
    for run in heading.runs:
        run.font.color.rgb = DARK_GRAY

    # Avatar info box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')

    p = cell.paragraphs[0]
    p.add_run('[Photo Placeholder]\n').bold = True
    p.add_run('Jasmine "Jazz" Carter\n').bold = True
    p.add_run('High School Senior • Basketball • Oakland, CA\n')
    p.add_run('12K Instagram • 8K TikTok')

    doc.add_paragraph()

    # The Story
    p = doc.add_paragraph()
    run = p.add_run('The Situation: ')
    run.bold = True
    p.add_run(
        'Senior year, Jazz\'s highlight reel goes viral. Within a week, she has 15 DMs from brands '
        'wanting to pay her for posts. She\'s excited—but also confused.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Challenge: ')
    run.bold = True
    p.add_run(
        'Is this even legal in California? Will accepting a deal affect her Stanford recruitment? '
        'Her mom is skeptical. Her coach says "be careful." But no one has actual answers.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Discovery: ')
    run.bold = True
    p.add_run(
        'Her school\'s athletic director introduces ChatNIL as a required educational tool for any '
        'athlete considering NIL activities.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Journey: ')
    run.bold = True
    p.add_run(
        'Day 1, the AI Coach asks about her goals—not her follower count. By Week 1, she learns '
        'California allows HS NIL with restrictions. Week 2, she completes the Identity pillar and '
        'understands her personal brand. Week 3, her mom approves consent after seeing it\'s educational. '
        'By Month 1, she\'s earned her first badge and knows what a legitimate deal looks like.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Outcome: ')
    run.bold = True
    p.add_run(
        'By graduation, Jazz knows the difference between a real opportunity and a scam. She turns down '
        'two sketchy offers. When she gets to Stanford, she\'s ready—and her compliance officer is impressed.'
    )

    # Pull quote
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"ChatNIL taught me what questions to ask before I even knew what questions to ask."')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = CHATNIL_ORANGE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('— Jasmine Carter')

    # Solution callout
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F97316')
    p = cell.paragraphs[0]
    run = p.add_run('How ChatNIL Helped: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run2 = p.add_run('Discovery conversation taught state rules, 4-pillar education prepared her for college NIL, parent consent kept her family involved and protected.')
    run2.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_page_break()

    # ==================== DARIUS'S STORY ====================
    heading = doc.add_heading("Darius's Story: \"The $25,000 Red Flag\"", 2)
    for run in heading.runs:
        run.font.color.rgb = DARK_GRAY

    # Avatar info box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')

    p = cell.paragraphs[0]
    p.add_run('[Photo Placeholder]\n').bold = True
    p.add_run('Darius "D-Money" Johnson\n').bold = True
    p.add_run('College Junior • Basketball • NC State University\n')
    p.add_run('85K Instagram Followers')

    doc.add_paragraph()

    # The Story
    p = doc.add_paragraph()
    run = p.add_run('The Situation: ')
    run.bold = True
    p.add_run(
        'D-Money is having a breakout season. A "sports marketing firm" offers him $25,000 for '
        '"brand ambassador" work. Easy money, right?'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Challenge: ')
    run.bold = True
    p.add_run(
        'Something feels off. The company name sounds like a booster collective. His teammate got '
        'suspended last year for a similar deal. But $25,000 is life-changing money.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Discovery: ')
    run.bold = True
    p.add_run(
        'NC State\'s compliance office requires all deals validated through ChatNIL before signing.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Journey: ')
    run.bold = True
    p.add_run(
        'Darius enters the deal details into the validator. The system returns a RED score: 42/100. '
        'The breakdown shows: FMV inflated 200%, booster-connected flag, vague deliverables. '
        'The AI explains why this screams "pay-for-play." Darius declines the deal.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Outcome: ')
    run.bold = True
    p.add_run(
        'Two weeks later, the "marketing firm" is exposed as a booster collective. Three athletes at '
        'rival schools lose eligibility. Darius finds a legitimate apparel deal for $8,000 that scores '
        'GREEN (88/100). He stays eligible and stays smart.'
    )

    # Pull quote
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"That RED score saved my career. I almost threw away everything for $25K."')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = CHATNIL_ORANGE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('— Darius Johnson')

    # Solution callout
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F97316')
    p = cell.paragraphs[0]
    run = p.add_run('How ChatNIL Helped: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run2 = p.add_run('6-dimension scoring flagged the deal as RED (booster-connected, inflated FMV). AI explained the risks. Darius declined and found a legitimate deal instead.')
    run2.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_page_break()

    # ==================== MICHELLE'S STORY ====================
    heading = doc.add_heading("Michelle's Story: \"I Finally Understand What My Daughter Is Doing\"", 2)
    for run in heading.runs:
        run.font.color.rgb = DARK_GRAY

    # Avatar info box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')

    p = cell.paragraphs[0]
    p.add_run('[Photo Placeholder]\n').bold = True
    p.add_run('Michelle Carter\n').bold = True
    p.add_run('Parent • Registered Nurse • Oakland, CA\n')
    p.add_run('Mother of Jasmine Carter')

    doc.add_paragraph()

    # The Story
    p = doc.add_paragraph()
    run = p.add_run('The Situation: ')
    run.bold = True
    p.add_run(
        'Jasmine asks permission to join "some NIL platform." Michelle\'s first thought: '
        '"What is NIL and why does my daughter need it?"'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Challenge: ')
    run.bold = True
    p.add_run(
        'She Googles NIL and finds horror stories—kids signing bad contracts, losing eligibility, '
        'getting scammed. She wants to say no, but doesn\'t want to hold Jasmine back.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Discovery: ')
    run.bold = True
    p.add_run(
        'The ChatNIL consent email explains exactly what the platform does and doesn\'t do. '
        'It\'s education, not a marketplace. No one is trying to sell her daughter to brands.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Journey: ')
    run.bold = True
    p.add_run(
        'Michelle reads the consent explanation and sees it\'s not connecting her daughter to brands. '
        'She creates a parent account and approves consent. Each week, she checks the dashboard and sees '
        'Jasmine earning badges. She gets a notification when Jasmine completes the Money pillar. '
        'She realizes her daughter now understands taxes better than most adults.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Outcome: ')
    run.bold = True
    p.add_run(
        'Michelle goes from skeptic to advocate. She tells other parents at Jasmine\'s games about ChatNIL. '
        '"It\'s the only platform that put my daughter\'s education first."'
    )

    # Pull quote
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"I went from \'What is NIL?\' to recommending ChatNIL to every parent I know."')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = CHATNIL_ORANGE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('— Michelle Carter')

    # Solution callout
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F97316')
    p = cell.paragraphs[0]
    run = p.add_run('How ChatNIL Helped: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run2 = p.add_run('Consent flow explained the platform clearly. Parent dashboard provided visibility without control. Activity feed showed education happening, not exploitation.')
    run2.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_page_break()

    # ==================== ANGELA'S STORY ====================
    heading = doc.add_heading("Angela's Story: \"Zero Violations in Year One\"", 2)
    for run in heading.runs:
        run.font.color.rgb = DARK_GRAY

    # Avatar info box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')

    p = cell.paragraphs[0]
    p.add_run('[Photo Placeholder]\n').bold = True
    p.add_run('Angela Washington, J.D.\n').bold = True
    p.add_run('Compliance Officer • Atlantic Coast University\n')
    p.add_run('D1 • 650 Athletes • 22 Sports')

    doc.add_paragraph()

    # The Story
    p = doc.add_paragraph()
    run = p.add_run('The Situation: ')
    run.bold = True
    p.add_run(
        'New NCAA rules, new state laws, and 650 athletes who all think they\'re the next NIL millionaire. '
        'Angela\'s inbox is drowning.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Challenge: ')
    run.bold = True
    p.add_run(
        'Her 4-person staff can\'t manually review every deal. Last year, another school missed a booster '
        'deal and got hit with a $2M penalty. She can\'t let that happen here.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Discovery: ')
    run.bold = True
    p.add_run(
        'Angela evaluates ChatNIL\'s compliance tools. The 6-dimension scoring system speaks her language. '
        'The audit trail is exactly what NCAA investigators ask for.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Journey: ')
    run.bold = True
    p.add_run(
        'Angela onboards all 650 athletes over two weeks. The dashboard immediately shows 12 athletes in '
        'RED status. She investigates: 8 are booster-connected deals, 4 have FMV issues. Athletes fix or '
        'decline the deals before signing. She exports NCAA-compliant reports with one click.'
    )

    p = doc.add_paragraph()
    run = p.add_run('The Outcome: ')
    run.bold = True
    p.add_run(
        'Year-end audit comes. Angela has documentation for every deal, every override, every decision. '
        'Zero violations. The AD asks her to present ChatNIL to the athletic conference. '
        '"This is how compliance should work."'
    )

    # Pull quote
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF7ED')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"ChatNIL gave me my weekends back. I\'m not chasing athletes for paperwork anymore."')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = CHATNIL_ORANGE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('— Angela Washington, J.D.')

    # Solution callout
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F97316')
    p = cell.paragraphs[0]
    run = p.add_run('How ChatNIL Helped: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run2 = p.add_run('Athletes self-validate deals. Real-time dashboard surfaces problems. Audit trail provides NCAA-ready documentation. Compliance at scale without additional staff.')
    run2.font.color.rgb = RGBColor(255, 255, 255)

    return doc

def main():
    # Open existing document
    input_path = '/Users/verrelbricejr./ChatNIL.io/docs/ChatNIL_Platform_Overview.docx'
    output_path = '/Users/verrelbricejr./ChatNIL.io/docs/ChatNIL_Platform_Overview.docx'

    print(f'Opening {input_path}...')
    doc = Document(input_path)

    print('Adding Customer Stories section...')
    add_customer_stories(doc)

    print(f'Saving to {output_path}...')
    doc.save(output_path)

    print('Done! Customer Stories section added.')
    return output_path

if __name__ == '__main__':
    main()
